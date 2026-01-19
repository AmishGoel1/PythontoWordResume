"""
YAML to DOCX Resume Conversion Module.

This module provides data models, rendering classes, and utilities for converting
YAML-structured resume data into professionally formatted Word documents (.docx).
It uses Pydantic for data validation and python-docx for document generation.

Key Components:
    Data Models (Pydantic):
        - skillsclass: Skill category and specific skill
        - workclass: Work experience entry with title, company, dates, and bullet points
        - projectclass: Project entry with name and description points
        - certificateclass: Certificate with name and issuer
        - credentialclass: Academic credential with name and details
        - educationclass: Educational institution with credentials
        - Resume: Main resume model containing all sections
        - ContactInfo: Contact information with validated email and URLs
    
    Rendering Classes (Abstract Base):
        - SectionRenderBaseClass: Abstract base class for section renderers
        - ProfessionalSummaryRenderer: Renders professional summary section
        - SkillRenderer: Renders categorized skills section
        - WorkExperienceRenderer: Renders work experience with bullet points
        - EducationRenderer: Renders education and certifications
        - ProjectRenderer: Renders projects with descriptions
    
    Utilities:
        - TextStyle: Dataclass for text formatting configuration
        - spacing: Function to set paragraph spacing
        - paragraph_formatting: Function to format and add paragraphs
        - formattingstyles: Dictionary of predefined text styles
        - renderermap: Mapping of section names to renderer instances

Dependencies:
    - pydantic: For data validation and parsing
    - python-docx: For Word document generation
"""
from abc import ABC, abstractmethod
from dataclasses import dataclass
from enum import Enum
from typing import List

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from pydantic import BaseModel, EmailStr, HttpUrl


class skillsclass(BaseModel):
    """
    Model representing a skill entry with category and skill name.
    
    Attributes:
        category: The category or type of skill (e.g., "Programming Languages", "Tools").
        skill: Comma-separated or descriptive list of skills in this category.
    """
    category: str
    skill: str

class workclass(BaseModel):
    """
    Model representing a work experience entry.
    
    Attributes:
        Title: Job title or position held.
        Company: Name of the company or organization.
        Date: Date range of employment (e.g., "Jan 2020 - Dec 2022").
        Points: List of bullet points describing responsibilities and achievements.
    """
    Title: str
    Company: str
    Date: str
    Points: List[str]

class projectclass(BaseModel):
    """
    Model representing a project entry.
    
    Attributes:
        Name: Name or title of the project.
        Points: List of bullet points describing the project details and accomplishments.
    """
    Name: str
    Points: List[str]

class certificateclass(BaseModel):
    """
    Model representing a certificate or certification.
    
    Attributes:
        Name: Name of the certificate or certification.
        Issuer: Organization or institution that issued the certificate.
    """
    Name: str
    Issuer: str

class credentialclass(BaseModel):
    """
    Model representing an academic credential (degree, diploma, etc.).
    
    Attributes:
        Name: Name of the credential (e.g., "Bachelor of Science in Computer Science").
        Points: List of additional details like GPA, honors, relevant coursework, etc.
    """
    Name: str
    Points: List[str]

class educationclass(BaseModel):
    """
    Model representing an educational institution and associated credentials.
    
    Attributes:
        Name: Name of the educational institution (e.g., "Stanford University").
        Credential: List of credentials obtained from this institution.
    """
    Name: str
    Credential: List[credentialclass]

class Resume(BaseModel):
    """
    Main resume model containing all resume sections.
    
    This is the top-level model that encapsulates all resume content
    including professional summary, skills, education, work experience,
    certificates, and projects.
    
    Attributes:
        summary: Professional summary or objective statement.
        skills: List of categorized skills.
        education: List of educational institutions and credentials.
        certificates: List of certifications obtained.
        work: List of work experience entries.
        projects: List of project entries.
    """
    summary: str
    skills: List[skillsclass]
    education: List[educationclass]
    certificates: List[certificateclass]
    work: List[workclass]
    projects: List[projectclass]

class SectionRenderBaseClass(ABC):
    """
    Abstract base class for rendering a resume section.
    
    All section renderers must inherit from this class and implement
    the render method to define how their specific section should be
    added to the Word document.
    
    Methods:
        render: Abstract method to render section content to a document.
    """
    @abstractmethod
    def render(self, doc: Document, data: Resume):
        """
        Render this section's content to the document.
        
        Args:
            doc: The Word document to which content should be added.
            data: Resume data containing all section information.
        """
        pass
class Sections(Enum):
    """
    Enumeration of resume section names.
    
    Defines standardized section names used throughout the resume builder
    for mapping section types to their corresponding renderers.
    """
    PROFESSIONAL_SUMMARY = 'Professional Summary'
    SKILLS = 'Skills'
    EXPERIENCE = 'Experience'
    EDUCATON = 'Education'
    PROJECTS = 'Projects'

def spacing(paragraph: Paragraph, space_before: Pt, space_after: Pt):
    """
    Set spacing before and after a paragraph.
    
    Args:
        paragraph: The paragraph object to modify.
        space_before: Space to add before the paragraph (in points).
        space_after: Space to add after the paragraph (in points).
    """
    paragraph.paragraph_format.space_after = space_after
    paragraph.paragraph_format.space_before = space_before
@dataclass
class TextStyle:
    """
    Configuration for text styling in the document.
    
    Defines formatting properties that can be applied to text runs
    in the Word document, including font properties.
    
    Attributes:
        bold: Whether text should be bold (default: False).
        font_name: Font family name (default: 'Times New Roman').
        font_size: Font size in points (default: 12pt).
    
    Methods:
        apply: Apply this style to a text run.
    """
    bold: bool = False
    font_name: str = 'Times New Roman'
    font_size: Pt = Pt(12)

    def apply(self, run: Run):
        """
        Apply this text style to a document run.
        
        Args:
            run: The text run object to which styling should be applied.
        """
        run.font.bold = self.bold
        run.font.name = self.font_name
        run.font.size = self.font_size

formattingstyles = {
    'Name': TextStyle(font_size=Pt(18)),
    'Links': TextStyle(font_size=Pt(14)),
    'Point': TextStyle(),
    'PointHeading': TextStyle(bold=True),
    'SectionHeading': TextStyle(bold=True, font_size=Pt(13))
}

class ProfessionalSummaryRenderer(SectionRenderBaseClass):
    """
    Renderer for the professional summary section.
    
    Adds the professional summary text as a paragraph with standard
    point formatting and minimal spacing.
    """
    def render(self, doc: Document, data: Resume):
        """
        Render professional summary to the document.
        
        Args:
            doc: Word document to add content to.
            data: Resume data containing the summary text.
        """
        text = doc.add_paragraph(f"{data.summary}\n")
        formattingstyles['Point'].apply(text.runs[0])   
        spacing(text, Pt(0), Pt(0))
        
class SkillRenderer(SectionRenderBaseClass):
    """
    Renderer for the skills section.
    
    Displays skills organized by category, with category names in bold
    followed by the skills in regular text.
    """
    def render(self, doc: Document, data: Resume):
        """
        Render skills section to the document.
        
        Each skill category is displayed as "Category: skills" with
        the category name in bold and skills in regular font.
        
        Args:
            doc: Word document to add content to.
            data: Resume data containing skills list.
        """
        for skillcategory in data.skills:
            categoryname = doc.add_paragraph(f"{skillcategory.category}: ")
            formattingstyles['PointHeading'].apply(categoryname.runs[0])
            valuerun = categoryname.add_run(f"{skillcategory.skill}")
            formattingstyles['Point'].apply(valuerun)
            categoryname.paragraph_format.space_after = Pt(3)
        doc.add_paragraph()

class WorkExperienceRenderer(SectionRenderBaseClass):
    """
    Renderer for the work experience section.
    
    Displays work experience entries with title, company, and dates on one line,
    followed by bullet points for responsibilities and achievements.
    """
    def render(self, doc: Document, data: Resume):
        """
        Render work experience section to the document.
        
        Each work entry includes a header line with position title, company,
        and dates, followed by bulleted achievement points.
        
        Args:
            doc: Word document to add content to.
            data: Resume data containing work experience list.
        """
        for work in data.work:
            doc.add_paragraph(f"{work.Title}, {work.Company} {work.Date}")
            for point in work.Points:
                currentworkpoint = doc.add_paragraph(f"{point}", style='List Bullet')
                formattingstyles['Point'].apply(currentworkpoint.runs[0])
                currentworkpoint.paragraph_format.space_after = Pt(3)
        doc.add_paragraph()

class EducationRenderer(SectionRenderBaseClass):
    """
    Renderer for the education section.
    
    Displays educational institutions with their credentials and details,
    followed by any certifications. Each institution shows its name,
    credentials (e.g., degrees), and additional details as bullet points.
    Certificates are listed separately with their issuers.
    """
    def render(self, doc: Document, data: Resume):
        """
        Render education and certifications section to the document.
        
        First renders educational institutions with their credentials and details,
        then renders certificates with their issuers.
        
        Args:
            doc: Word document to add content to.
            data: Resume data containing education and certificates lists.
        """
        for university in data.education:
                universityname = doc.add_paragraph(university.Name)
                formattingstyles['PointHeading'].apply(universityname.runs[0])
                for credential in university.Credential:
                        credentialname = doc.add_paragraph(credential.Name)
                        credentialname.paragraph_format.space_after = Pt(0)
                        formattingstyles['PointHeading'].apply(credentialname.runs[0])
                        for point in credential.Points:
                            currentpoint = doc.add_paragraph(point, style='List Bullet')
                            formattingstyles['Point'].apply(currentpoint.runs[0])
                            currentpoint.paragraph_format.space_after = Pt(0)
        doc.add_paragraph()
        for certificate in data.certificates:
            certificatename = doc.add_paragraph(f"{certificate.Name}: {certificate.Issuer}", style='List Bullet')
            formattingstyles['Point'].apply(certificatename.runs[0])
            certificatename.paragraph_format.space_after = Pt(2)
class ProjectRenderer(SectionRenderBaseClass):
    """
    Renderer for the projects section.
    
    Displays project entries with project names in bold followed by
    bullet points describing project details and accomplishments.
    Applies custom spacing for visual hierarchy.
    """
    def render(self, doc: Document, data: Resume):
        """
        Render projects section to the document.
        
        Each project shows its name in bold followed by descriptive
        bullet points. Special spacing is applied to the first project.
        
        Args:
            doc: Word document to add content to.
            data: Resume data containing projects list.
        """
        for project in data.projects:
            projectname = doc.add_paragraph(project.Name)
            formattingstyles['PointHeading'].apply(projectname.runs[0])
            if project.Name == 'Project 1 Name':
                projectname.paragraph_format.space_before = Pt(0)
            else: 
                projectname.paragraph_format.space_before = Pt(9)
            for point in project.Points:
                currentpoint = doc.add_paragraph(f"{point}", style='List Bullet')
                formattingstyles['Point'].apply(currentpoint.runs[0])
                spacing(currentpoint, Pt(15), Pt(3))
class ContactInfo(BaseModel):
    """
    Model for contact information with validation.
    
    Validates and stores contact information including name, email,
    and social media URLs.
    
    Attributes:
        name: Full name of the resume owner.
        email: Email address (validated by Pydantic's EmailStr).
        github: GitHub profile URL (validated by Pydantic's HttpUrl).
        linkedin: LinkedIn profile URL (validated by Pydantic's HttpUrl).
    """
    name: str
    email: EmailStr
    github: HttpUrl
    linkedin: HttpUrl

# contact = ContactInfo(
#     name = 'AMISH GOEL',
#     email = 'Jobs@lets.workwithamish.me',
#     github = "https://google.com", # type: ignore
#     linkedin = "https://github.com/AmishGoel1" # type: ignore
# )

def paragraph_formatting(doc: Document, paragraph_text: str,style: TextStyle, alignment = WD_ALIGN_PARAGRAPH.CENTER):
    """
    Create and format a paragraph with specified text and style.
    
    Adds a new paragraph to the document with the given text, applies
    the specified text style, and sets paragraph alignment.
    
    Args:
        doc: Word document to add the paragraph to.
        paragraph_text: Text content for the paragraph.
        style: TextStyle object defining formatting properties.
        alignment: Paragraph alignment (default: center-aligned).
    
    Returns:
        Paragraph: The created and formatted paragraph object.
    """
    para = doc.add_paragraph(paragraph_text)
    TextStyle().apply(para.runs[0])
    para.alignment = alignment

    return para
renderermap: dict[str, SectionRenderBaseClass] = {
    Sections.PROFESSIONAL_SUMMARY.value: ProfessionalSummaryRenderer(),
    Sections.SKILLS.value: SkillRenderer(),
    Sections.EDUCATON.value: EducationRenderer(),
    Sections.EXPERIENCE.value: WorkExperienceRenderer(),
    Sections.PROJECTS.value: ProjectRenderer()
}

