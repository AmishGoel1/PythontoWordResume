from docx import Document as doc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from pydantic import BaseModel
from typing import List
import yaml
from dataclasses import dataclass
from enum import Enum
from abc import ABC, abstractmethod
from docx.document import Document

try:
    with open("points.yaml", "r") as file:
        data = yaml.safe_load(file)
    print("Successfully read prompt text file.")
except FileNotFoundError:
    print("Prompt text file not found. Please make sure the file exists")
except PermissionError:
	print("Permission denied when accessing the file. Please make sure you can access the file")
except IOError as e:
    print(f"An I/O error occured {e}")
class skillsclass(BaseModel):
    category: str
    skill: str

class workclass(BaseModel):
    Title: str
    Company: str
    Date: str
    Points: List[str]

class projectclass(BaseModel):
    Name: str
    Points: List[str]

class certificateclass(BaseModel):
    Name: str
    Issuer: str

class credentialclass(BaseModel):
    Name: str
    Points: List[str]

class educationclass(BaseModel):
    Name: str
    Credential: List[credentialclass]

class Resume(BaseModel):
    summary: str
    skills: List[skillsclass]
    education: List[educationclass]
    certificates: List[certificateclass]
    work: List[workclass]
    projects: List[projectclass]

class SectionRenderBaseClass(ABC):
    """Abstract Class for Rendering a section
    
    Keyword arguments:
    ABC class -- Inheriting from ABC abstract class
    Return: nothing
    """
    @abstractmethod
    def render(self, doc: Document, data: Resume):
        pass
class Sections(Enum):
    PROFESSIONAL_SUMMARY = 'Professional Summary'
    SKILLS = 'Skills'
    EXPERIENCE = 'Experience'
    EDUCATON = 'Education'
    PROJECTS = 'Projects'

def spacing(paragraph: Paragraph, space_before: Pt, space_after: Pt):
    paragraph.paragraph_format.space_after = space_after
    paragraph.paragraph_format.space_before = space_before
@dataclass
class TextStyle:
    bold: bool = False
    font_name: str = 'Times New Roman'
    font_size: Pt = Pt(12)

    def apply(self, run: Run):
        run.font.bold = self.bold
        run.font.name = self.font_name
        run.font.size = self.font_size

formattingstyles = {
    'Name': TextStyle(font_size=Pt(18)),
    'Links': TextStyle(font_size=Pt(14)),
    'Point': TextStyle(),
    'PointHeading': TextStyle(bold=True),
    'SectionHeading': TextStyle(bold=True, font_size=Pt(13))
}

class ProfessionalSummaryRenderer(SectionRenderBaseClass):
    def render(self, doc: Document, data: Resume):
        text = doc.add_paragraph(f"{data.summary}\n")
        formattingstyles['Point'].apply(text.runs[0])   
        spacing(text, Pt(0), Pt(0))
        
class SkillRenderer(SectionRenderBaseClass):
    def render(self, doc: Document, data: Resume):
        for skillcategory in data.skills:
            categoryname = doc.add_paragraph(f"{skillcategory.category}: ")
            formattingstyles['PointHeading'].apply(categoryname.runs[0])
            valuerun = categoryname.add_run(f"{skillcategory.skill}")
            formattingstyles['Point'].apply(valuerun)
            categoryname.paragraph_format.space_after = Pt(3)
        doc.add_paragraph()

class WorkExperienceRenderer(SectionRenderBaseClass):
    def render(self, doc: Document, data: Resume):
        for work in resume.work:
            doc.add_paragraph(f"{work.Title}, {work.Company} {work.Date}")
            for point in work.Points:
                currentworkpoint = doc.add_paragraph(f"{point}", style='List Bullet')
                formattingstyles['Point'].apply(currentworkpoint.runs[0])
                currentworkpoint.paragraph_format.space_after = Pt(3)
        doc.add_paragraph()

class EducationRenderer(SectionRenderBaseClass):
    def render(self, doc: Document, data: Resume):
        for university in resume.education:
                universityname = doc.add_paragraph(university.Name)
                formattingstyles['PointHeading'].apply(universityname.runs[0])
                for credential in university.Credential:
                        credentialname = doc.add_paragraph(credential.Name)
                        credentialname.paragraph_format.space_after = Pt(0)
                        formattingstyles['PointHeading'].apply(credentialname.runs[0])
                        for point in credential.Points:
                            currentpoint = doc.add_paragraph(point, style='List Bullet')
                            formattingstyles['Point'].apply(currentpoint.runs[0])
                            currentpoint.paragraph_format.space_after = Pt(0)
        doc.add_paragraph()
        for certificate in resume.certificates:
            certificatename = doc.add_paragraph(f"{certificate.Name}: {certificate.Issuer}", style='List Bullet')
            formattingstyles['Point'].apply(certificatename.runs[0])
            certificatename.paragraph_format.space_after = Pt(2)
    
class ProjectRenderer(SectionRenderBaseClass):
    def render(self, doc: Document, data: Resume):
        for project in resume.projects:
            projectname = doc.add_paragraph(project.Name)
            formattingstyles['PointHeading'].apply(projectname.runs[0])
            if project.Name == 'Project 1 Name':
                projectname.paragraph_format.space_before = Pt(0)
            else: 
                projectname.paragraph_format.space_before = Pt(9)
            for point in project.Points:
                currentpoint = doc.add_paragraph(f"{point}", style='List Bullet')
                formattingstyles['Point'].apply(currentpoint.runs[0])
                spacing(currentpoint, Pt(15), Pt(3))

resume = Resume.model_validate(data['resume'][0])
name = 'AMISH GOEL'
email = 'Jobs@lets.workwithamish.me'
linkedinLink = 'https://www.linkedin.com/in/amish-goyal-096066b4/'
github = 'https://github.com/AmishGoel1'
initialdoc = doc()

renderermap: dict[str, SectionRenderBaseClass] = {
    Sections.PROFESSIONAL_SUMMARY.value: ProfessionalSummaryRenderer(),
    Sections.SKILLS.value: SkillRenderer(),
    Sections.EDUCATON.value: EducationRenderer(),
    Sections.EXPERIENCE.value: WorkExperienceRenderer(),
    Sections.PROJECTS.value: ProjectRenderer()
}

nametext = initialdoc.add_paragraph('AMISH GOEL')
formattingstyles['Name'].apply(nametext.runs[0])
nametext.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
nametext.paragraph_format.space_after = Pt(0)

linkspara = initialdoc.add_paragraph(f"{email} | LinkedIn | GitHub")
formattingstyles['Links'].apply(linkspara.runs[0])
linkspara.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

for section in data['resume_sections']:
    heading = initialdoc.add_paragraph(section['type'])
    formattingstyles['SectionHeading'].apply(heading.runs[0])
    renderermap[section['type']].render(initialdoc, resume)

section = initialdoc.sections[0]
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
section.top_margin = Inches(0.437)
section.bottom_margin = Inches(0.287)

initialdoc.save('sapmle1.docx')

